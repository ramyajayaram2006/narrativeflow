import io
import re
import json
import string
import streamlit as st
from collections import Counter
from styles import workspace_style
from database import (save_story, add_character, load_characters, delete_character,
                      update_character, add_scene, load_scenes, delete_scene, update_scene)
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

LLM_BACKEND  = "ollama"
OLLAMA_MODEL = "llama3.2"

PLOT_STAGES = [
    ("beginning",      "🌱 Beginning"),
    ("rising_action",  "📈 Rising Action"),
    ("climax",         "⚡ Climax"),
    ("falling_action", "📉 Falling Action"),
    ("resolution",     "🌅 Resolution"),
]

_STOPWORDS = {
    "the","a","an","and","or","but","in","on","at","to","for","of","with","by","from",
    "as","is","was","are","were","be","been","being","have","has","had","do","does","did",
    "will","would","could","should","may","might","shall","can","i","you","he","she","it",
    "we","they","me","him","her","us","them","my","your","his","its","our","their",
    "this","that","these","those","which","who","what","when","where","how","if","then",
    "so","not","no","up","out","about","into","than","more","just","also","said","like",
    "there","here","all","some","one","two","back","into","very","over","her","his"
}


# ── Core helpers ───────────────────────────────────────────────────────────────
def _get_story():
    for s in st.session_state.stories:
        if s["id"] == st.session_state.current_story:
            return s
    return None

def _word_count(story):
    return sum(len(m["content"].split()) for m in story.get("messages", []))

def _all_prose(story):
    return " ".join(m["content"] for m in story.get("messages", [])
                    if not m["content"].startswith("🌿") and not m["content"].startswith("✍️"))

def _sentences(text):
    return [s.strip() for s in re.split(r'[.!?]+', text) if len(s.strip()) > 4]

def _is_incomplete(text):
    t = text.strip()
    if not t or t[-1] in ".!?": return False
    last = t.split()[-1].lower().rstrip(",:;")
    return last in {"and","but","or","with","from","to","as","while","when","the","a","an","of","in"}


# ── Story-Only Classification — 7-Instruction System ──────────────────────────
#
# Instruction 1: Classify every input as story or non-story
# Instruction 2: Story input → respond normally
# Instruction 3: Non-story input → redirect, never answer
# Instruction 4: Use one of three redirect tones (soft/firm/guiding)
# Instruction 5: Outside knowledge FOR a story = allowed (mixed requests)
# Instruction 6: Block math, coding, health, general knowledge, random facts
# Instruction 7: Never break the boundary — all non-story inputs redirected

# ── Instruction 1: Story keyword vocabulary ───────────────────────────────────
# If any of these appear in the input → classify as STORY (Instruction 5 handled here)
_STORY_KEYWORDS = {
    # Core narrative
    "plot","subplot","story","narrative","tale","fiction","novel","chapter",
    "arc","story arc","structure","act","scene","scenes","sequence","event",
    # Characters
    "character","characters","protagonist","antagonist","villain","hero",
    "heroine","narrator","mentor","sidekick","foil","backstory","motivation",
    "character arc","character development","persona","voice",
    # World & setting
    "worldbuilding","world","setting","lore","universe","realm","location",
    "atmosphere","environment","time period","era","dystopia","utopia",
    # Craft
    "dialogue","monologue","theme","themes","tone","mood","pacing","tension",
    "conflict","resolution","climax","foreshadowing","flashback","imagery",
    "metaphor","symbolism","point of view","pov","perspective","style",
    "writing style","genre","draft","revision","edit","editing","revise",
    "prose","description","exposition","show dont tell","subtext","irony",
    # Actions
    "write","writing","create","creating","develop","build","craft","describe",
    "rewrite","improve","strengthen","deepen","outline","plan","brainstorm",
    # Story elements
    "opening","beginning","ending","twist","reveal","stakes","motivation",
    "emotion","fear","love","betrayal","redemption","journey","quest",
    "timeline","flashforward","prologue","epilogue","chapter",
    # Genres & forms
    "fantasy","sci-fi","romance","mystery","horror","thriller","adventure",
    "historical fiction","literary fiction","short story","screenplay",
}

# ── Instruction 6: Explicitly blocked request types ───────────────────────────
_BLOCKED_PREFIXES = [
    # Math
    "calculate","solve","what is","what are","how much","how many",
    "add ","subtract","multiply","divide","compute",
    # Coding
    "code","program","script","function","debug","fix my code","write code",
    "write a program","how to code","how do i code","python","javascript",
    "html","css","sql","api","algorithm",
    # General knowledge
    "who is","who was","when did","when was","where is","where was",
    "tell me about","explain to me","give me facts","what happened",
    "history of","define ","what does","what's the meaning",
    # Health / advice
    "health","medical","symptom","diagnosis","medicine","diet","exercise",
    "calories","nutrition","doctor","treatment",
    # Factual queries
    "weather","temperature","news","score","price","stock","currency",
    "translate","convert","how to make","recipe","install","download",
    # Catch-all homework / off-topic help
    "help me with","help me find","help me understand","help me learn",
    "help me fix my","help me solve","help me calculate",
    # Contractions that bypass Q_STARTERS
    "what's","whats","who's","whos","where's","wheres","how's","hows",
    "isn't","aren't","wasn't","weren't","doesn't","don't","didn't",
]

# Question starters that signal non-story intent (when no story keyword present)
_Q_STARTERS = {
    "what","who","where","when","why","how","is","are","was","were",
    "do","does","did","can","could","would","should","will","have","has","had",
    "which","whose","give","list","find","search","lookup","calculate","solve",
}

# ── Redirect message variants (Instruction 4) ─────────────────────────────────
import random as _random

_REDIRECTS_SOFT = (
    "I'm designed to support story creation. Share a story idea and I'll help you build it. ✍️",
    "I'm here to help with storytelling and creative writing. What are you creating — a plot, a character, or a scene?",
)
_REDIRECTS_FIRM = (
    "I focus only on storytelling. Let's return to your story — what would you like to work on?",
    "I'm here to help with stories only. Tell me about your plot, character, or world. 🌿",
)
_REDIRECTS_GUIDING = (
    "I can't help with that, but I can help shape a story. What would you like to create? **Plot · Character · Scene · Dialogue · Worldbuilding**",
    "That's outside my storytelling focus. Let's keep things creative — do you have a character, scene, or plot idea to explore? ✍️",
)

def _get_redirect():
    pool = _REDIRECTS_SOFT + _REDIRECTS_FIRM + _REDIRECTS_GUIDING
    return _random.choice(pool)


# ── Instruction 1 classifier ──────────────────────────────────────────────────
def _classify_input(text):
    """
    Classifies input as 'story' or 'redirect' using the 7-instruction system.

    Instruction 5 — Mixed requests:
        Outside knowledge in story context is allowed.
        e.g. "explain fear psychology for my horror story"
        → contains 'story' and 'horror' → classified as STORY.

    Instruction 6 — Blocked types:
        Math, coding, general knowledge, health, random facts → redirect.
    """
    t = text.strip()
    if not t:
        return 'story'

    t_lower = t.lower()
    words = set(t_lower.replace(',','').replace('.','').replace('?','').split())

    # Instruction 5 — story keyword anywhere = story (even in questions)
    # e.g. "explain fear for my horror story" → has 'horror', 'story' → STORY
    if words & _STORY_KEYWORDS:
        return 'story'

    # Instruction 6 — blocked request types
    for prefix in _BLOCKED_PREFIXES:
        if t_lower.startswith(prefix):
            return 'redirect'

    # General question with no story context
    if t.endswith("?"):
        return 'redirect'
    if t.split()[0].lower().rstrip(",:;") in _Q_STARTERS:
        return 'redirect'

    # Instruction 7 — default: short prose → story (benefit of the doubt)
    return 'story'

def _is_non_story(text):
    return _classify_input(text) == 'redirect'


# ── Prompt builder ─────────────────────────────────────────────────────────────
def _build_prompt(user_input, story, mode, characters=None, scenes=None):
    genre = story["genre"]
    tone  = story["tone"].lower()
    writing_style = story.get("writing_style","").strip()
    history = story.get("messages", [])
    context = "\n".join(f"{m['role'].capitalize()}: {m['content']}"
                        for m in history[-10:]
                        if not m["content"].startswith("🌿"))

    style_block = f"\nWriting voice: {writing_style}\n" if writing_style else ""

    char_block = ""
    if characters:
        lines = []
        for c in characters:
            line = f"  - {c['name']} ({c['role']}): {c['description']}"
            if c.get("speaking_style"):
                line += f" Their dialogue style: {c['speaking_style']}."
            lines.append(line)
        char_block = "\nCharacters in this story:\n" + "\n".join(lines) + "\n"

    scene_block = ""
    if scenes:
        sc = scenes[-1]
        scene_block = (f"\nCurrent scene: '{sc['title']}' — "
                       f"Location: {sc['location']} — Purpose: {sc['purpose']}\n")

    base = (f"You are a creative writing assistant. Genre: {genre}. Tone: {tone}."
            f"{style_block}{char_block}{scene_block}"
            f"\nStory so far:\n{context}\n\n")

    if mode == "continue":
        return base + (f"Complete this unfinished sentence in 1-2 sentences. "
                       f"Do NOT repeat the user's text:\n{user_input}")
    elif mode == "paragraph":
        return base + (f"The user just wrote: {user_input}\n\n"
                       f"Write the next paragraph (3-5 sentences). "
                       f"Keep all characters consistent with their descriptions.")
    else:
        return _build_prompt(user_input, story,
                             "continue" if _is_incomplete(user_input) else "paragraph",
                             characters, scenes)

def _char_suggest_prompt(story, characters):
    excerpt = " ".join(m["content"] for m in story.get("messages",[])[-4:])[:600]
    existing_names = [c["name"] for c in characters] if characters else []
    existing_str = ", ".join(existing_names) if existing_names else "none yet"
    avoid = (f"\nIMPORTANT: Do NOT suggest any of these existing characters: {existing_str}. "
             f"The new character must have a completely different name and role.")  if existing_names else ""
    return (
        f"You are helping a writer build a character for their {story['genre']} story "
        f"(tone: {story['tone'].lower()}).\n"
        f"Story excerpt: {excerpt}\nExisting characters: {existing_str}{avoid}\n\n"
        f"Suggest ONE brand new character not already in the story. "
        f"Respond in EXACTLY this format, nothing else:\n"
        f"Name: [character name]\n"
        f"Role: [protagonist/antagonist/supporting/mentor/narrator]\n"
        f"Description: [2 sentences about personality and appearance]\n"
        f"Speaking style: [1 sentence about how they talk — pace, vocabulary, verbal habits]"
    )


# ── AI backends ────────────────────────────────────────────────────────────────
def _stream_ollama(prompt):
    """Generator — yields text tokens for real-time streaming"""
    import requests
    try:
        with requests.post(
            "http://localhost:11434/api/generate",
            json={"model": OLLAMA_MODEL, "prompt": prompt, "stream": True,
                  "options": {"temperature": 0.85, "num_predict": 280}},
            stream=True, timeout=90
        ) as r:
            if r.status_code == 404:
                yield f"[Model not found — run: ollama pull {OLLAMA_MODEL}]"
                return
            for line in r.iter_lines():
                if line:
                    try:
                        chunk = json.loads(line)
                        token = chunk.get("response","")
                        if token: yield token
                        if chunk.get("done"): break
                    except: continue
    except Exception as e:
        if "Connection refused" in str(e):
            yield "⚠️ Ollama not running. Open a terminal and run: ollama serve"
        else:
            yield f"[Error: {e}]"

def _call_full(prompt, max_tokens=120):
    """Non-streaming — for redirects and suggestions"""
    import requests
    try:
        r = requests.post("http://localhost:11434/api/generate",
                          json={"model": OLLAMA_MODEL, "prompt": prompt, "stream": False,
                                "options": {"temperature": 0.78, "num_predict": max_tokens}},
                          timeout=30)
        return r.json().get("response","").strip() if r.status_code == 200 else ""
    except: return ""


# ── Story Health ───────────────────────────────────────────────────────────────
def _reading_level(text):
    words = text.split()
    sents = _sentences(text)
    if not words or not sents: return "—"
    syllables = sum(max(1, len(re.findall(r'[aeiouAEIOU]', w))) for w in words)
    fk = 0.39*(len(words)/len(sents)) + 11.8*(syllables/len(words)) - 15.59
    grade = max(1, min(16, round(fk)))
    if grade <= 5:  return f"Simple (Gr. {grade})"
    if grade <= 8:  return f"Easy (Gr. {grade})"
    if grade <= 11: return f"Medium (Gr. {grade})"
    return f"Advanced (Gr. {grade})"

def _story_health(story):
    prose = _all_prose(story)
    words = prose.split()
    sents = _sentences(prose)
    total = len(words)
    clean = [w.lower().strip(string.punctuation) for w in words]
    freq  = Counter(w for w in clean if w and w not in _STOPWORDS and len(w) > 2)
    return {
        "total_words":    total,
        "total_sents":    len(sents),
        "avg_sent_len":   round(total / len(sents), 1) if sents else 0,
        "reading_label":  _reading_level(prose),
        "top_words":      freq.most_common(8),
        "word_goal":      story.get("word_goal", 0),
        "messages":       len(story.get("messages",[])),
    }

def _consistency_check(story, characters):
    prose  = _all_prose(story)
    words  = prose.split()
    issues = []
    total  = len(words) or 1

    clean = [w.lower().strip(string.punctuation) for w in words]
    freq  = Counter(w for w in clean if w and w not in _STOPWORDS and len(w) > 3)
    for word, count in freq.most_common(15):
        if count/total*100 > 3.0 and count >= 5:
            issues.append({"type":"overused","text":f'"{word}"',
                           "detail":f"used {count}x ({count/total*100:.1f}% of words)"})

    trigrams = [" ".join(words[i:i+3]).lower() for i in range(len(words)-2)]
    for phrase, count in Counter(trigrams).most_common(5):
        if count >= 3 and not any(sw in phrase.split()[0] for sw in list(_STOPWORDS)[:8]):
            issues.append({"type":"phrase","text":f'"{phrase}"',"detail":f"repeated {count}x"})

    long_s = [s for s in _sentences(prose) if len(s.split()) > 35]
    if long_s:
        issues.append({"type":"sentence","text":"Long sentences",
                       "detail":f"{len(long_s)} sentence(s) over 35 words"})

    if characters and len(characters) > 1:
        char_counts = {c["name"]: prose.lower().count(c["name"].lower()) for c in characters}
        total_m = sum(char_counts.values()) or 1
        for name, count in char_counts.items():
            if count/total_m*100 < 5 and total_m > 20:
                issues.append({"type":"character","text":name,
                               "detail":f"only mentioned {count}x — may feel absent"})
    return issues[:8]


# ── Word export ────────────────────────────────────────────────────────────────
def _export_docx(story, characters, scenes):
    doc = Document()
    t = doc.add_heading(story["title"], 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m = doc.add_paragraph()
    m.add_run(f"Genre: {story['genre']}  |  Tone: {story['tone']}  |  Words: {_word_count(story)}").italic = True
    m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    if story.get("writing_style"):
        doc.add_heading("Writing Voice", 1)
        doc.add_paragraph(story["writing_style"])
        doc.add_paragraph()

    if characters:
        doc.add_heading("Characters", 1)
        for c in characters:
            p = doc.add_paragraph()
            p.add_run(f"{c['name']}").bold = True
            p.add_run(f"  ({c['role']})")
            doc.add_paragraph(c["description"])
            if c.get("speaking_style"):
                sp = doc.add_paragraph()
                sp.add_run("Speaking style: ").bold = True
                sp.add_run(c["speaking_style"])
        doc.add_paragraph()

    arc = story.get("plot_arc", {})
    completed = [label for key, label in PLOT_STAGES if arc.get(key)]
    if completed:
        doc.add_heading("Plot Arc", 1)
        doc.add_paragraph(" → ".join(completed))
        doc.add_paragraph()

    if scenes:
        doc.add_heading("Scene Overview", 1)
        for i, sc in enumerate(scenes, 1):
            p = doc.add_paragraph()
            p.add_run(f"Scene {i}: {sc['title']}").bold = True
            if sc["location"]: doc.add_paragraph(f"Location: {sc['location']}")
            if sc["purpose"]:  doc.add_paragraph(f"Purpose: {sc['purpose']}")
            if sc["characters"]: doc.add_paragraph(f"Characters: {', '.join(sc['characters'])}")
        doc.add_paragraph()

    doc.add_heading("Story", 1)
    doc.add_paragraph("─" * 50)
    prose = " ".join(m["content"] for m in story.get("messages",[])
                     if m["role"] == "assistant" and not m["content"].startswith("🌿"))
    doc.add_paragraph(prose.strip() if prose else "No content yet.")
    doc.add_paragraph()
    f = doc.add_paragraph()
    f.add_run("Generated by NarrativeFlow").italic = True
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── Sidebar panels ─────────────────────────────────────────────────────────────
def _sidebar_settings(story):
    wc   = _word_count(story)
    goal = story.get("word_goal", 0)
    pct  = min(100, int(wc / goal * 100)) if goal else None

    st.markdown(f"""
        <div class='wc-box'>
            <div class='wc-number'>{wc:,}</div>
            <div class='wc-label'>Words Written</div>
        </div>
    """, unsafe_allow_html=True)

    if pct is not None:
        color = "#4ade80" if pct < 90 else "#86efac"
        st.markdown(f"""
            <div style='margin:4px 0 8px;'>
                <div style='display:flex;justify-content:space-between;font-size:0.72rem;color:#4b7a56;margin-bottom:3px;'>
                    <span>Goal</span><span>{pct}% of {goal:,}</span>
                </div>
                <div style='background:rgba(74,222,128,0.1);border-radius:4px;height:6px;'>
                    <div style='width:{pct}%;background:{color};height:6px;border-radius:4px;'></div>
                </div>
            </div>
        """, unsafe_allow_html=True)

    st.markdown(f"<div class='wc-sub'>{len(story.get('messages',[]))} messages</div>", unsafe_allow_html=True)
    st.markdown("---")

    new_title = st.text_input("Story Title", value=story["title"], key="ws_title")
    if new_title != story["title"]: story["title"] = new_title

    genres = ["Fantasy","Sci-Fi","Mystery","Romance","Horror","Adventure","Thriller","Historical"]
    story["genre"] = st.selectbox("Genre", genres,
        index=genres.index(story.get("genre","Fantasy")) if story.get("genre") in genres else 0,
        key="ws_genre")
    tones = ["Dark","Light","Emotional","Humorous","Serious","Suspenseful","Whimsical"]
    story["tone"] = st.selectbox("Tone", tones,
        index=tones.index(story.get("tone","Light")) if story.get("tone") in tones else 1,
        key="ws_tone")

    st.markdown("<div style='margin-top:10px;font-size:0.78rem;color:#4b7a56;font-weight:600;text-transform:uppercase;letter-spacing:0.08em;'>✍️ Writing Voice</div>", unsafe_allow_html=True)
    new_style = st.text_area("_", value=story.get("writing_style",""),
                              placeholder="e.g. lyrical, slow-burn, first-person past tense, rich sensory details",
                              height=75, key="ws_style", label_visibility="collapsed")
    if new_style != story.get("writing_style",""):
        story["writing_style"] = new_style
        save_story(st.session_state.username, story)

    goal_input = st.number_input("🎯 Word Goal", min_value=0, max_value=200000,
                                  value=int(story.get("word_goal",0)), step=500, key="ws_goal")
    if goal_input != story.get("word_goal",0):
        story["word_goal"] = goal_input
        save_story(st.session_state.username, story)

    st.markdown(f"<div style='color:#4b7a56;font-size:0.75rem;margin-top:6px;'>🤖 AI: <code style='color:#4ade80;'>{LLM_BACKEND}</code></div>", unsafe_allow_html=True)
    st.markdown("---")

    chars  = load_characters(st.session_state.username, story["id"])
    scenes = load_scenes(st.session_state.username, story["id"])
    if story.get("messages"):
        st.download_button("📥 Download Story",
            data=_export_docx(story, chars, scenes),
            file_name=f"{story['title'].replace(' ','_')}_NarrativeFlow.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True)

    if st.button("🗑️ Clear Chat", use_container_width=True):
        story["messages"] = []
        save_story(st.session_state.username, story)
        st.session_state.show_cowrite_options = False
        st.rerun()
    if st.button("🚪 Logout", use_container_width=True):
        st.session_state.authenticated = False
        st.session_state.current_view = "auth"
        st.rerun()


def _sidebar_characters(story):
    username   = st.session_state.username
    # FIX: Always load fresh from DB — never use cached list
    characters = load_characters(username, story["id"])
    roles_list = ["protagonist","antagonist","supporting","mentor","narrator"]

    if characters:
        for c in characters:
            voice_html = f"<div class='char-voice'>🗣 {c['speaking_style']}</div>" if c.get("speaking_style") else ""
            st.markdown(f"""
                <div class='char-card'>
                    <div class='char-name'>{c['name']}</div>
                    <div class='char-role'>{c['role']}</div>
                    <div class='char-desc'>{c['description']}</div>
                    {voice_html}
                </div>
            """, unsafe_allow_html=True)

            # Edit + Delete buttons side by side
            col_e, col_d = st.columns(2)
            with col_e:
                if st.button("✏️ Edit", key=f"editchar_{c['id']}", use_container_width=True):
                    # Toggle: click again to close
                    if st.session_state.get("editing_char") == c["id"]:
                        st.session_state.pop("editing_char", None)
                    else:
                        st.session_state["editing_char"] = c["id"]
                    st.rerun()
            with col_d:
                if st.button("🗑️ Delete", key=f"delchar_{c['id']}", use_container_width=True):
                    delete_character(c["id"])
                    st.session_state.pop("editing_char", None)
                    st.rerun()

            # Inline edit form — only shown for the selected character
            if st.session_state.get("editing_char") == c["id"]:
                st.markdown("<div style='background:rgba(74,222,128,0.05);border:1px solid rgba(74,222,128,0.2);border-radius:8px;padding:10px;margin:6px 0;'>", unsafe_allow_html=True)
                e_name  = st.text_input("Name",           value=c["name"],  key=f"ename_{c['id']}")
                e_role  = st.selectbox("Role", roles_list,
                            index=roles_list.index(c["role"]) if c["role"] in roles_list else 2,
                            key=f"erole_{c['id']}")
                e_desc  = st.text_area("Description",     value=c["description"], height=80, key=f"edesc_{c['id']}")
                e_voice = st.text_input("Speaking Style",  value=c.get("speaking_style",""), key=f"evoice_{c['id']}")
                sv1, sv2 = st.columns(2)
                with sv1:
                    if st.button("💾 Save", key=f"savechar_{c['id']}", use_container_width=True):
                        update_character(c["id"], e_name, e_role, e_desc, e_voice)
                        st.session_state.pop("editing_char", None)
                        st.rerun()
                with sv2:
                    if st.button("✖ Cancel", key=f"cancelchar_{c['id']}", use_container_width=True):
                        st.session_state.pop("editing_char", None)
                        st.rerun()
                st.markdown("</div>", unsafe_allow_html=True)
    else:
        st.markdown("<div style='color:#4b7a56;font-size:0.82rem;text-align:center;padding:12px;'>No characters yet.</div>", unsafe_allow_html=True)

    st.markdown("---")

    # AI Suggest — always queries fresh list, always clears previous suggestion
    if st.button("✨ AI Suggest", use_container_width=True, type="primary"):
        for k in ["char_suggest","char_name","char_role","char_desc","char_voice"]:
            st.session_state.pop(k, None)
        fresh_chars = load_characters(username, story["id"])
        with st.spinner("Thinking of a new character..."):
            raw = _call_full(_char_suggest_prompt(story, fresh_chars), max_tokens=160)
        s = {"name":"","role":"supporting","description":"","speaking_style":""}
        for line in raw.splitlines():
            ll = line.lower()
            if ll.startswith("name:"):        s["name"]          = line.split(":",1)[1].strip()
            elif ll.startswith("role:"):
                r = line.split(":",1)[1].strip().lower()
                s["role"] = r if r in roles_list else "supporting"
            elif ll.startswith("description:"): s["description"] = line.split(":",1)[1].strip()
            elif "speaking" in ll and ":" in ll: s["speaking_style"] = line.split(":",1)[1].strip()
        st.session_state["char_suggest"] = s
        st.rerun()

    sug = st.session_state.get("char_suggest", {})
    with st.expander("➕ Add Character", expanded=bool(sug)):
        c_name  = st.text_input("Name",  value=sug.get("name",""),  key="char_name")
        c_role  = st.selectbox("Role", roles_list,
                    index=roles_list.index(sug.get("role","supporting")), key="char_role")
        c_desc  = st.text_area("Description", value=sug.get("description",""),
                    height=70, key="char_desc", placeholder="Personality, appearance, backstory...")
        c_voice = st.text_input("Speaking Style", value=sug.get("speaking_style",""),
                    key="char_voice", placeholder="e.g. formal and slow, speaks in riddles...")
        if st.button("💾 Save Character", use_container_width=True):
            if c_name.strip():
                add_character(username, story["id"], c_name, c_role, c_desc, c_voice)
                for k in ["char_suggest","char_name","char_role","char_desc","char_voice"]:
                    st.session_state.pop(k, None)
                st.rerun()


def _sidebar_scenes(story):
    username   = st.session_state.username
    # FIX: Always load fresh from DB
    scenes     = load_scenes(username, story["id"])
    characters = load_characters(username, story["id"])
    # FIX: Always reflect current character list — never stale
    char_names = [c["name"] for c in characters]

    if scenes:
        for sc in scenes:
            chars_html = f"<div class='scene-meta'>👥 {', '.join(sc['characters'])}</div>" if sc["characters"] else ""
            st.markdown(f"""
                <div class='scene-card'>
                    <div class='scene-title'>Scene {sc['order']}: {sc['title']}</div>
                    <div class='scene-meta'>📍 {sc['location'] or '—'} · 🎯 {sc['purpose'] or '—'}</div>
                    {chars_html}
                </div>
            """, unsafe_allow_html=True)

            se1, se2 = st.columns(2)
            with se1:
                if st.button("✏️ Edit", key=f"editscene_{sc['id']}", use_container_width=True):
                    if st.session_state.get("editing_scene") == sc["id"]:
                        st.session_state.pop("editing_scene", None)
                    else:
                        st.session_state["editing_scene"] = sc["id"]
                    st.rerun()
            with se2:
                if st.button("🗑️ Delete", key=f"delscene_{sc['id']}", use_container_width=True):
                    delete_scene(sc["id"])
                    st.session_state.pop("editing_scene", None)
                    st.rerun()

            # Inline edit form for this scene
            if st.session_state.get("editing_scene") == sc["id"]:
                st.markdown("<div style='background:rgba(74,222,128,0.05);border:1px solid rgba(74,222,128,0.2);border-radius:8px;padding:10px;margin:6px 0;'>", unsafe_allow_html=True)
                es_title    = st.text_input("Title",    value=sc["title"],    key=f"estitle_{sc['id']}")
                es_location = st.text_input("Location", value=sc["location"], key=f"esloc_{sc['id']}")
                es_purpose  = st.text_input("Purpose",  value=sc["purpose"],  key=f"espurp_{sc['id']}")
                # FIX: Character multiselect uses CURRENT character list, pre-selects existing
                current_in_scene = [ch for ch in sc["characters"] if ch in char_names]
                es_chars = st.multiselect("Characters", char_names,
                             default=current_in_scene, key=f"eschars_{sc['id']}") if char_names else []
                ss1, ss2 = st.columns(2)
                with ss1:
                    if st.button("💾 Save", key=f"savescene_{sc['id']}", use_container_width=True):
                        update_scene(sc["id"], es_title, es_location, es_purpose, es_chars)
                        st.session_state.pop("editing_scene", None)
                        st.rerun()
                with ss2:
                    if st.button("✖ Cancel", key=f"cancelscene_{sc['id']}", use_container_width=True):
                        st.session_state.pop("editing_scene", None)
                        st.rerun()
                st.markdown("</div>", unsafe_allow_html=True)
    else:
        st.markdown("<div style='color:#4b7a56;font-size:0.82rem;text-align:center;padding:12px;'>No scenes yet.</div>", unsafe_allow_html=True)

    st.markdown("---")
    with st.expander("➕ Add Scene"):
        s_title    = st.text_input("Title",    key="s_title")
        s_location = st.text_input("Location", key="s_loc")
        s_purpose  = st.text_input("Purpose",  key="s_purpose",
                                   placeholder="e.g. Introduce conflict")
        # FIX: Multiselect always shows current characters — empty list if none added yet
        s_chars = st.multiselect("Characters in this scene", char_names, key="s_chars") if char_names else []
        if not char_names:
            st.caption("Add characters in the Cast tab first to assign them to scenes.")
        if st.button("💾 Save Scene", use_container_width=True):
            if s_title.strip():
                add_scene(username, story["id"], s_title, s_location, s_purpose, s_chars)
                for k in ["s_title","s_loc","s_purpose","s_chars"]:
                    st.session_state.pop(k, None)
                st.rerun()


def _sidebar_arc(story):
    arc     = story.get("plot_arc", {})
    changed = False
    st.markdown("<div style='color:#4b7a56;font-size:0.75rem;margin-bottom:10px;'>Tick off stages as you write them.</div>", unsafe_allow_html=True)
    for key, label in PLOT_STAGES:
        val = st.checkbox(label, value=arc.get(key,False), key=f"arc_{story['id']}_{key}")
        if val != arc.get(key,False): arc[key] = val; changed = True
    if changed:
        story["plot_arc"] = arc
        save_story(st.session_state.username, story)

    done = sum(1 for k,_ in PLOT_STAGES if arc.get(k))
    pct  = int(done/len(PLOT_STAGES)*100)
    st.markdown(f"""
        <div style='margin-top:12px;background:rgba(74,222,128,0.06);border:1px solid rgba(74,222,128,0.15);
            border-radius:10px;padding:10px;text-align:center;'>
            <div style='font-size:1.4rem;font-weight:700;color:#4ade80;'>{pct}%</div>
            <div style='font-size:0.7rem;color:#4b7a56;text-transform:uppercase;letter-spacing:0.1em;'>Arc Complete</div>
            <div style='font-size:0.78rem;color:#86efac;margin-top:4px;'>{done}/{len(PLOT_STAGES)} stages</div>
        </div>
    """, unsafe_allow_html=True)

    scenes = load_scenes(st.session_state.username, story["id"])
    if scenes:
        st.markdown("---")
        st.markdown("<div style='color:#4b7a56;font-size:0.72rem;text-transform:uppercase;letter-spacing:0.1em;margin-bottom:8px;'>📅 Timeline</div>", unsafe_allow_html=True)
        for sc in scenes:
            st.markdown(f"""
                <div style='display:flex;gap:8px;margin:6px 0;'>
                    <div style='width:6px;height:6px;border-radius:50%;background:#4ade80;margin-top:5px;flex-shrink:0;'></div>
                    <div>
                        <div style='font-size:0.8rem;color:#f0fdf4;font-weight:500;'>Scene {sc['order']}: {sc['title']}</div>
                        <div style='font-size:0.72rem;color:#4b7a56;'>{sc.get('location','')}</div>
                    </div>
                </div>
            """, unsafe_allow_html=True)


def _sidebar_health(story, characters):
    if not story.get("messages"):
        st.markdown("<div style='color:#4b7a56;font-size:0.82rem;text-align:center;padding:20px;'>Write some story first.</div>", unsafe_allow_html=True)
        return

    h      = _story_health(story)
    issues = _consistency_check(story, characters)

    st.markdown(f"""
        <div class='health-grid'>
            <div class='health-stat'>
                <div class='hs-num'>{h['avg_sent_len']}</div>
                <div class='hs-lbl'>Avg Sentence<br>Length</div>
            </div>
            <div class='health-stat'>
                <div class='hs-num' style='font-size:0.85rem;'>{h['reading_label']}</div>
                <div class='hs-lbl'>Reading<br>Level</div>
            </div>
        </div>
    """, unsafe_allow_html=True)
    st.markdown("---")

    if h["top_words"]:
        st.markdown("<div style='font-size:0.72rem;color:#4b7a56;text-transform:uppercase;letter-spacing:0.08em;margin-bottom:8px;font-weight:600;'>📊 Word Frequency</div>", unsafe_allow_html=True)
        max_count = h["top_words"][0][1] if h["top_words"] else 1
        for word, count in h["top_words"]:
            total_w = max(h["total_words"],1)
            density = count/total_w*100
            bar_w   = int(count/max_count*100)
            color   = "#e11d48" if density > 3.5 else "#4ade80"
            st.markdown(f"""
                <div style='display:flex;align-items:center;gap:8px;margin:4px 0;'>
                    <div style='width:65px;font-size:0.78rem;color:#f0fdf4;font-weight:500;overflow:hidden;white-space:nowrap;'>{word}</div>
                    <div style='flex:1;background:rgba(74,222,128,0.08);border-radius:3px;height:5px;'>
                        <div style='width:{bar_w}%;background:{color};height:5px;border-radius:3px;'></div>
                    </div>
                    <div style='width:28px;font-size:0.72rem;color:#4b7a56;text-align:right;'>{count}x</div>
                </div>
            """, unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("<div style='font-size:0.72rem;color:#4b7a56;text-transform:uppercase;letter-spacing:0.08em;margin-bottom:8px;font-weight:600;'>🔍 Consistency</div>", unsafe_allow_html=True)
    if not issues:
        st.markdown("<div style='color:#4ade80;font-size:0.82rem;text-align:center;padding:8px;'>✅ No issues found!</div>", unsafe_allow_html=True)
    else:
        icons  = {"overused":"🔁","phrase":"💬","sentence":"📏","character":"👤"}
        colors = {"overused":"#fca5a5","phrase":"#fcd34d","sentence":"#93c5fd","character":"#c4b5fd"}
        for issue in issues:
            ic  = icons.get(issue["type"],"⚠️")
            clr = colors.get(issue["type"],"#86efac")
            st.markdown(f"""
                <div class='issue-card' style='border-left-color:{clr};'>
                    <div style='font-size:0.8rem;color:#f0fdf4;font-weight:600;'>{ic} {issue['text']}</div>
                    <div style='font-size:0.72rem;color:#4b7a56;margin-top:2px;'>{issue['detail']}</div>
                </div>
            """, unsafe_allow_html=True)


# ── Main workspace ─────────────────────────────────────────────────────────────

def _force_sidebar_open():
    """Inject JS to force sidebar open if Streamlit collapsed it"""
    import streamlit.components.v1 as components
    components.html("""
        <script>
        // Force sidebar open after short delay (wait for Streamlit to render)
        setTimeout(function() {
            const sidebar = window.parent.document.querySelector('[data-testid="stSidebar"]');
            const collapseBtn = window.parent.document.querySelector('[data-testid="collapsedControl"]');
            if (sidebar) {
                sidebar.style.display = 'block';
                sidebar.style.visibility = 'visible';
                sidebar.style.transform = 'none';
            }
            // Click the expand button if sidebar is collapsed
            if (collapseBtn && collapseBtn.style.display !== 'none') {
                collapseBtn.click();
            }
            // Also try the chevron button
            const chevrons = window.parent.document.querySelectorAll('button[kind="header"]');
            chevrons.forEach(btn => {
                const rect = btn.getBoundingClientRect();
                if (rect.left < 30) btn.click();  // Only click left-edge buttons (sidebar toggle)
            });
        }, 300);
        </script>
    """, height=0)

def show_workspace():
    workspace_style()
    _force_sidebar_open()
    story = _get_story()
    if not story:
        st.error("Story not found.")
        if st.button("← Back"):
            st.session_state.current_view = "dashboard"; st.rerun()
        return

    username   = st.session_state.username

    # ── Top navigation bar ────────────────────────────────────────────────────
    nav1, nav2 = st.columns([1, 8])
    with nav1:
        if st.button("← Dashboard", key="topnav_back", use_container_width=True):
            save_story(username, story)
            st.session_state.current_view = "dashboard"
            st.session_state.show_cowrite_options = False
            st.rerun()

    # ── Sidebar ────────────────────────────────────────────────────────────────
    with st.sidebar:
        st.markdown(f"### 👤 {username}")
        st.markdown("---")
        t1, t2, t3, t4, t5 = st.tabs(["⚙️ Set", "👥 Cast", "🎬 Scenes", "📈 Arc", "💊 Health"])
        with t1: _sidebar_settings(story)
        with t2: _sidebar_characters(story)
        with t3: _sidebar_scenes(story)
        with t4: _sidebar_arc(story)
        with t5:
            _chars = load_characters(username, story["id"])
            _sidebar_health(story, _chars)

    # ── Fresh loads for header, prompt, and export ───────────────────────────
    characters = load_characters(username, story["id"])
    scenes     = load_scenes(username, story["id"])

    # ── Header ─────────────────────────────────────────────────────────────────
    arc  = story.get("plot_arc", {})
    done = sum(1 for k,_ in PLOT_STAGES if arc.get(k))
    voice_badge = f"<span class='ws-badge'>✍️ {story['writing_style'][:22]}…</span>" if story.get("writing_style") else ""
    arc_badge   = f"<span class='ws-badge ws-badge-green'>📈 Arc {done}/5</span>" if done else ""
    st.markdown(f"""
        <div style='margin-bottom:6px;'>
            <span style='font-family:"Playfair Display",serif;font-size:1.9rem;font-weight:700;color:#4ade80;'>
                📖 {story['title']}</span>
        </div>
        <div style='display:flex;gap:8px;margin-bottom:16px;flex-wrap:wrap;'>
            <span class='ws-badge'>📚 {story['genre']}</span>
            <span class='ws-badge'>🎭 {story['tone']}</span>
            <span class='ws-badge'>🤖 {LLM_BACKEND}</span>
            {arc_badge}{voice_badge}
        </div>
    """, unsafe_allow_html=True)

    # ── Chat history ────────────────────────────────────────────────────────────
    if story.get("messages"):
        for i, msg in enumerate(story["messages"]):
            if msg["role"] == "user":
                st.markdown(f"<div class='bubble-user'><div class='lbl'>🧑‍💻 You</div><div class='txt'>{msg['content']}</div></div>",
                            unsafe_allow_html=True)
            else:
                content = msg["content"]
                is_redirect = (
                    content.startswith("I'm here to help") or
                    content.startswith("I'm designed to") or
                    content.startswith("I focus only") or
                    content.startswith("I can't help with that") or
                    content.startswith("That's outside") or
                    (content.startswith("🌿") and len(content) < 400)
                )
                if is_redirect:
                    st.markdown(f"<div class='redirect-msg'>{content}</div>", unsafe_allow_html=True)
                else:
                    st.markdown(f"<div class='bubble-ai'><div class='lbl'>🌿 NarrativeFlow</div><div class='txt'>{content}</div></div>",
                                unsafe_allow_html=True)
                    if i == len(story["messages"]) - 1:
                        b1, b2 = st.columns(2)
                        with b1:
                            if st.button("🔄 Regenerate", key=f"regen_{i}", use_container_width=True):
                                prev   = story["messages"][i-1]["content"] if i > 0 else ""
                                prompt = _build_prompt(prev, story, "paragraph", characters, scenes)
                                slot   = st.empty()
                                tokens = []
                                for token in _stream_ollama(prompt):
                                    tokens.append(token)
                                    slot.markdown(f"<div class='bubble-ai'><div class='lbl'>🌿 NarrativeFlow</div><div class='txt'>{''.join(tokens)}</div></div>",
                                                  unsafe_allow_html=True)
                                story["messages"][i]["content"] = "".join(tokens)
                                save_story(username, story); st.rerun()
                        with b2:
                            if st.button("✍️ Continue This", key=f"cont_{i}", use_container_width=True):
                                prompt = _build_prompt(content, story, "paragraph", characters, scenes)
                                slot   = st.empty()
                                tokens = []
                                for token in _stream_ollama(prompt):
                                    tokens.append(token)
                                    slot.markdown(f"<div class='bubble-ai'><div class='lbl'>🌿 NarrativeFlow</div><div class='txt'>{''.join(tokens)}</div></div>",
                                                  unsafe_allow_html=True)
                                story["messages"].append({"role":"assistant","content":"".join(tokens)})
                                save_story(username, story); st.rerun()
    else:
        st.markdown("""
            <div style='text-align:center;padding:50px 20px;background:linear-gradient(135deg,#0f1e0f,#142014);
                border:1px solid rgba(74,222,128,0.12);border-radius:20px;margin:20px 0;'>
                <div style='font-size:2.5rem;margin-bottom:12px;'>🌿</div>
                <div style='color:#86efac;font-size:1.05rem;font-family:"Lora",serif;'>Begin your story below.</div>
                <div style='color:#4b7a56;font-size:0.85rem;margin-top:6px;'>Type a line or scene — the AI will stream a continuation.</div>
                <div style='color:#4b7a56;font-size:0.8rem;margin-top:4px;'>Add your writing voice in ⚙️ Settings to shape the style.</div>
            </div>
        """, unsafe_allow_html=True)

    # ── Input ───────────────────────────────────────────────────────────────────
    user_input = st.chat_input("Write your story here...")
    if user_input:
        story["messages"].append({"role":"user","content":user_input})
        save_story(username, story)

        if _is_non_story(user_input):
            # Instructions 3 & 4: redirect with a varied polite message, never answer
            story["messages"].append({"role":"assistant","content":_get_redirect()})
            save_story(username, story)
            st.session_state.show_cowrite_options = False
        else:
            st.session_state.show_cowrite_options = True
            st.session_state.pending_input = user_input
        st.rerun()

    # ── Mode buttons with live streaming ────────────────────────────────────────
    if st.session_state.get("show_cowrite_options"):
        st.markdown("<div class='mode-header'>🌿 How should I continue?</div>", unsafe_allow_html=True)
        m1, m2, m3 = st.columns(3)

        def _run_stream(mode_key):
            prompt = _build_prompt(st.session_state.pending_input, story, mode_key, characters, scenes)
            slot   = st.empty()
            tokens = []
            for token in _stream_ollama(prompt):
                tokens.append(token)
                slot.markdown(
                    f"<div class='bubble-ai'><div class='lbl'>🌿 NarrativeFlow</div>"
                    f"<div class='txt'>{''.join(tokens)}</div></div>",
                    unsafe_allow_html=True)
            story["messages"].append({"role":"assistant","content":"".join(tokens)})
            save_story(username, story)
            st.session_state.show_cowrite_options = False
            st.rerun()

        with m1:
            if st.button("✏️ Continue Sentence", use_container_width=True, type="primary"):
                _run_stream("continue")
        with m2:
            if st.button("📝 Add Paragraph", use_container_width=True, type="primary"):
                _run_stream("paragraph")
        with m3:
            if st.button("🎲 Smart Choice", use_container_width=True, type="secondary"):
                _run_stream("auto")
